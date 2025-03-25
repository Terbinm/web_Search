import os
from docx import Document


def setup_template_directory():
    """
    確保模板目錄存在並創建示範模板
    """
    # 獲取app目錄
    app_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    # 創建模板目錄
    template_dir = os.path.join(app_dir, 'static', 'templates')
    if not os.path.exists(template_dir):
        os.makedirs(template_dir)

    # 模板文件路徑
    template_path = os.path.join(template_dir, 'part_template.docx')

    # 如果模板文件不存在，創建默認模板
    if not os.path.exists(template_path):
        create_default_template(template_path)
        print(f"已創建默認模板: {template_path}")
    else:
        print(f"模板已存在: {template_path}")


def create_default_template(template_path):
    """
    創建默認的料號申編單模板
    """
    doc = Document()

    # 添加標題
    # 添加基本信息
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph('料號: {{ pn }}')
    doc.add_paragraph('英文品名: {{ english_name }}')
    doc.add_paragraph('中文品名: {{ chinese_name }}')
    doc.add_paragraph('美金單價: {{ price_usd }}')

    # 添加編號信息
    doc.add_heading('編號信息', level=1)
    doc.add_paragraph('單位會計編號: {{ accounting_number }}')
    doc.add_paragraph('品名代號: {{ item_code }}')
    doc.add_paragraph('撥發單位: {{ issuing_department }}')
    doc.add_paragraph('廠家代號: {{ vendor_code }}')
    doc.add_paragraph('參考號碼(P/N): {{ reference_number }}')
    doc.add_paragraph('P/N獲得程度: {{ pn_acquisition_level }}')
    doc.add_paragraph('P/N獲得來源: {{ pn_acquisition_source }}')

    # 添加規格信息
    doc.add_heading('規格信息', level=1)
    doc.add_paragraph('規格指示: {{ specification_indicator }}')
    doc.add_paragraph('單位包裝量: {{ packaging_quantity }}')
    doc.add_paragraph('存儲壽限: {{ storage_life }}')
    doc.add_paragraph('壽限處理: {{ storage_process }}')
    doc.add_paragraph('儲存型式: {{ storage_type }}')
    doc.add_paragraph('規格說明:{{ specification_description }}')

    # 添加艦艇與配置信息
    doc.add_heading('艦艇與配置信息', level=1)
    doc.add_paragraph('艦型: {{ ship_category }}')
    doc.add_paragraph('CID/NO.: {{ configuration_id }}')
    doc.add_paragraph('型式: {{ model_id }}')
    doc.add_paragraph('品名: {{ item_name }}')
    doc.add_paragraph('裝置數: {{ installation_number }}')
    doc.add_paragraph('位置: {{ location }}')

    # 添加其他信息
    doc.add_heading('其他信息', level=1)
    # 創建表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    # 填充表格
    cells = table.rows[0].cells
    cells[0].text = '機密性代號'
    cells[1].text = '{{ classification }}'

    cells = table.rows[1].cells
    cells[0].text = '消耗性代號'
    cells[1].text = '{{ consumability }}'

    cells = table.rows[2].cells
    cells[0].text = '修理能量'
    cells[1].text = '{{ repair_capability }}'

    cells = table.rows[3].cells
    cells[0].text = '製造能量'
    cells[1].text = '{{ manufacturing_capability }}'

    cells = table.rows[4].cells
    cells[0].text = '來源代碼'
    cells[1].text = '{{ source }}'

    cells = table.rows[5].cells
    cells[0].text = '系統代號'
    cells[1].text = '{{ system }}'

    cells = table.rows[6].cells
    cells[0].text = '檔別代號'
    cells[1].text = '{{ category }}'

    cells = table.rows[7].cells
    cells[0].text = '文件參考來源'
    cells[1].text = '{{ pn_acquisition_source }}'

    # 添加簽章區域
    doc.add_heading('簽章區域', level=1)
    doc.add_paragraph('申請單位: {{ application_unit }}')
    doc.add_paragraph('申請日期: {{ application_date }}')
    doc.add_paragraph('申請單位簽章: {{ application_unit_signature }}')
    doc.add_paragraph('審核單位簽章: {{ review_unit_signature }}')
    doc.add_paragraph('NC建檔會辦單簽章: {{ nc_file_unit_signature }}')
    # 保存文檔
    doc.save(template_path)


if __name__ == "__main__":
    setup_template_directory()
"""
DOCX生成器工具 - 用於將表單數據填充到Word模板中
"""
import os
from docx import Document
from jinja2 import Template


def replace_placeholders_jinja2(doc_path, output_path, variables):
    """
    使用jinja2語法替換Word文檔中的佔位符

    參數:
    doc_path (str): 輸入Word文檔路徑
    output_path (str): 輸出Word文檔路徑
    variables (dict): 變數字典，鍵為變數名，值為替換內容
    """
    try:
        doc = Document(doc_path)

        # 遍歷文檔中的所有段落
        for paragraph in doc.paragraphs:
            if any(x in paragraph.text for x in ['{{', '{%', '{#']):
                # 使用jinja2處理文本
                template = Template(paragraph.text)
                try:
                    paragraph.text = template.render(**variables)
                except Exception as e:
                    print(f"處理段落時出錯: {e}")
                    print(f"問題段落: {paragraph.text}")

        # 遍歷表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(x in paragraph.text for x in ['{{', '{%', '{#']):
                            # 使用jinja2處理文本
                            template = Template(paragraph.text)
                            try:
                                paragraph.text = template.render(**variables)
                            except Exception as e:
                                print(f"處理表格單元格時出錯: {e}")
                                print(f"問題段落: {paragraph.text}")

        # 保存文檔
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"生成DOCX文件時出錯: {e}")
        return False


def generate_part_docx(form_data, output_path):
    """
    根據表單數據生成料號申編單DOCX文件

    參數:
    form_data (dict): 表單數據
    output_path (str): 輸出文件路徑

    返回:
    bool: 是否成功生成文件
    """
    try:
        # 獲取模板文件路徑
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(current_dir, '..', 'static', 'templates', 'part_template.docx')

        # 確保模板目錄存在
        template_dir = os.path.dirname(template_path)
        if not os.path.exists(template_dir):
            os.makedirs(template_dir)

        # 如果模板文件不存在，創建一個基本模板
        if not os.path.exists(template_path):
            create_default_template(template_path)

        # 處理表單數據
        variables = {}
        for key, value in form_data.items():
            if value:  # 只處理非空值
                variables[key] = value

        # 執行替換並生成文件
        return replace_placeholders_jinja2(template_path, output_path, variables)
    except Exception as e:
        print(f"生成料號申編單出錯: {e}")
        return False


def create_default_template(template_path):
    """
    如果模板文件不存在，創建一個默認模板

    參數:
    template_path (str): 模板文件路徑
    """
    doc = Document()

    # 添加標題
    doc.add_heading('海軍料號申編單', 0)

    # 添加基本信息
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph('料號: {{ pn }}')
    doc.add_paragraph('英文品名: {{ english_name }}')
    doc.add_paragraph('中文品名: {{ chinese_name }}')
    doc.add_paragraph('美金單價: {{ price_usd }}')

    # 添加編號信息
    doc.add_heading('編號信息', level=1)
    doc.add_paragraph('單位會計編號: {{ accounting_number }}')
    doc.add_paragraph('品名代號: {{ item_code }}')
    doc.add_paragraph('撥發單位: {{ issuing_department }}')
    doc.add_paragraph('廠家代號: {{ vendor_code }}')
    doc.add_paragraph('參考號碼(P/N): {{ reference_number }}')
    doc.add_paragraph('P/N獲得程度: {{ pn_acquisition_level }}')
    doc.add_paragraph('P/N獲得來源: {{ pn_acquisition_source }}')

    # 添加規格信息
    doc.add_heading('規格信息', level=1)
    doc.add_paragraph('規格指示: {{ specification_indicator }}')
    doc.add_paragraph('單位包裝量: {{ packaging_quantity }}')
    doc.add_paragraph('存儲壽限: {{ storage_life }}')
    doc.add_paragraph('壽限處理: {{ storage_process }}')
    doc.add_paragraph('儲存型式: {{ storage_type }}')
    doc.add_paragraph('規格說明:{{ specification_description }}')

    # 添加艦艇與配置信息
    doc.add_heading('艦艇與配置信息', level=1)
    doc.add_paragraph('艦型: {{ ship_category }}')
    doc.add_paragraph('CID/NO.: {{ configuration_id }}')
    doc.add_paragraph('型式: {{ model_id }}')
    doc.add_paragraph('品名: {{ item_name }}')
    doc.add_paragraph('裝置數: {{ installation_number }}')
    doc.add_paragraph('位置: {{ location }}')

    # 添加其他信息
    doc.add_heading('其他信息', level=1)
    # 創建表格
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    # 填充表格
    cells = table.rows[0].cells
    cells[0].text = '機密性代號'
    cells[1].text = '{{ classification }}'

    cells = table.rows[1].cells
    cells[0].text = '消耗性代號'
    cells[1].text = '{{ consumability }}'

    cells = table.rows[2].cells
    cells[0].text = '修理能量'
    cells[1].text = '{{ repair_capability }}'

    cells = table.rows[3].cells
    cells[0].text = '製造能量'
    cells[1].text = '{{ manufacturing_capability }}'

    cells = table.rows[4].cells
    cells[0].text = '來源代碼'
    cells[1].text = '{{ source }}'

    cells = table.rows[5].cells
    cells[0].text = '系統代號'
    cells[1].text = '{{ system }}'

    cells = table.rows[6].cells
    cells[0].text = '檔別代號'
    cells[1].text = '{{ category }}'

    cells = table.rows[7].cells
    cells[0].text = '文件參考來源'
    cells[1].text = '{{ pn_acquisition_source }}'

    # 添加簽章區域
    doc.add_heading('簽章區域', level=1)
    doc.add_paragraph('申請單位: {{ application_unit }}')
    doc.add_paragraph('申請日期: {{ application_date }}')
    doc.add_paragraph('申請單位簽章: {{ application_unit_signature }}')
    doc.add_paragraph('審核單位簽章: {{ review_unit_signature }}')
    doc.add_paragraph('NC建檔會辦單簽章: {{ nc_file_unit_signature }}')

    # 保存文檔
    doc.save(template_path)